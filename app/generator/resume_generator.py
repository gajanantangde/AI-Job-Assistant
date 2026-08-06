from docx import Document


def generate_resume(resume_data, output_path):

    document = Document()

    document.add_heading(resume_data["full_name"], level=1)

    document.add_paragraph(f"Email: {resume_data['email']}")
    document.add_paragraph(f"Phone: {resume_data['phone']}")

    document.add_heading("Professional Summary", level=2)
    document.add_paragraph(resume_data["summary"])

    document.add_heading("Skills", level=2)

    skills = resume_data["skills"]

    if isinstance(skills, str):
        skills = [s.strip() for s in skills.split(",")]

    for skill in skills:
        document.add_paragraph(skill, style="List Bullet")

    document.add_heading("Projects", level=2)

    projects = resume_data["projects"]

    if projects:
        for project in projects.split("\n"):
            if project.strip():
                document.add_paragraph(project.strip(), style="List Bullet")

    document.save(output_path)

    return output_path